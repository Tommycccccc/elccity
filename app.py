import re
import io
import pandas as pd
import streamlit as st

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

st.set_page_config(page_title="ELC - City Directory Search", layout="wide")
st.title("ELC - City Directory Search")

# ---------- Styling ----------
st.markdown(
    """
    <style>
      div.stButton > button {
        width: 100%;
        height: 52px;
        border-radius: 10px;
        font-weight: 700;
        letter-spacing: .5px;
      }

      .addr-card {
        border: 1px solid rgba(255,255,255,.10);
        border-radius: 12px;
        padding: 14px 16px 10px 16px;
        background: rgba(255,255,255,.03);
        margin: 10px 0 18px 0;
      }

      .addr-header {
        font-size: 16px;
        font-weight: 800;
        margin: 0 0 10px 0;
        display: flex;
        align-items: center;
        flex-wrap: wrap;
        gap: 8px;
      }
      .addr-pill {
        display: inline-block;
        padding: 4px 10px;
        border-radius: 999px;
        background: rgba(255, 75, 75, .16);
        border: 1px solid rgba(255, 75, 75, .45);
        color: #ff4b4b;
        font-weight: 900;
        letter-spacing: .4px;
      }

      .neat-table {
        width: 100%;
        border-collapse: collapse;
        margin-top: 6px;
        table-layout: fixed;
      }
      .neat-table th, .neat-table td {
        padding: 10px 12px;
        border-bottom: 1px solid rgba(255,255,255,.08);
        vertical-align: top;
      }
      .neat-table th {
        text-align: left;
        font-size: 14px;
        opacity: .9;
      }
      .neat-table td:first-child, .neat-table th:first-child {
        width: 90px;
        text-align: center;
        font-variant-numeric: tabular-nums;
        white-space: nowrap;
      }
      .neat-table td:last-child, .neat-table th:last-child {
        word-break: break-word;
        overflow-wrap: anywhere;
      }

      .section-title { margin-top: 18px; }
    </style>
    """,
    unsafe_allow_html=True,
)

# ---------- Helpers ----------
def normalize_addr(addr: str) -> str:
    if addr is None:
        return ""
    s = str(addr).strip()
    s = re.sub(r"\s+", " ", s)
    return s

def parse_address_for_sort(addr: str) -> tuple[str, int, int, str]:
    """
    Sort key:
      (street_name_upper, house_number_int, unit_number_int, full_addr_upper)

    - street_name_upper: everything after the leading house number and optional unit (#)
    - house_number_int: leading number if present, else 0
    - unit_number_int: #N if present, else 0
    - full_addr_upper: stable tie-breaker
    """
    if addr is None:
        return ("", 0, 0, "")

    a = normalize_addr(str(addr))
    a_up = a.upper()

    # house number (leading digits)
    m_house = re.match(r"^\s*(\d+)", a_up)
    house = int(m_house.group(1)) if m_house else 0

    # unit number like "#3"
    m_unit = re.search(r"#\s*(\d+)", a_up)
    unit = int(m_unit.group(1)) if m_unit else 0

    # remove leading house number + space
    rest = re.sub(r"^\s*\d+\s*", "", a_up)

    # remove unit marker(s) like "#3" anywhere
    rest = re.sub(r"\s*#\s*\d+\s*", " ", rest)

    # normalize spaces -> street name key
    street = re.sub(r"\s+", " ", rest).strip()

    return (street, house, unit, a_up)

def find_and_combine_address_columns(df: pd.DataFrame) -> pd.DataFrame:
    """
    Find address column(s) and combine them into a single ADDRESS column.
    Handles:
    - Single ADDRESS column
    - ADDRESS1 + ADDRESS2 columns
    - Other address column variations
    """
    cols_upper = {col.upper(): col for col in df.columns}
    
    # Pattern 1: Single ADDRESS column exists
    if "ADDRESS" in cols_upper:
        df["ADDRESS"] = df[cols_upper["ADDRESS"]].apply(normalize_addr)
        return df
    
    # Pattern 2: ADDRESS1 and ADDRESS2 columns (combine them)
    if "ADDRESS1" in cols_upper and "ADDRESS2" in cols_upper:
        addr1_col = cols_upper["ADDRESS1"]
        addr2_col = cols_upper["ADDRESS2"]
        
        def combine_addresses(row):
            a1 = str(row[addr1_col]).strip() if pd.notna(row[addr1_col]) else ""
            a2 = str(row[addr2_col]).strip() if pd.notna(row[addr2_col]) else ""
            
            # Combine with a space, remove extra spaces
            combined = f"{a1} {a2}".strip()
            return normalize_addr(combined)
        
        df["ADDRESS"] = df.apply(combine_addresses, axis=1)
        st.info(f"✓ Combined {addr1_col} and {addr2_col} into ADDRESS column")
        return df
    
    # Pattern 3: Only ADDRESS1 exists (use it)
    if "ADDRESS1" in cols_upper:
        df["ADDRESS"] = df[cols_upper["ADDRESS1"]].apply(normalize_addr)
        st.info(f"✓ Using {cols_upper['ADDRESS1']} as ADDRESS column")
        return df
    
    # Pattern 4: Look for other common patterns
    address_patterns = [
        "STREET ADDRESS", "STREET_ADDRESS", "PROPERTY ADDRESS", 
        "PROPERTY_ADDRESS", "SITE ADDRESS", "LOCATION", "STREET", "ADDR"
    ]
    
    for pattern in address_patterns:
        if pattern in cols_upper:
            df["ADDRESS"] = df[cols_upper[pattern]].apply(normalize_addr)
            st.info(f"✓ Using {cols_upper[pattern]} as ADDRESS column")
            return df
    
    # No address column found
    return df

def find_listing_column(df: pd.DataFrame) -> pd.DataFrame:
    """
    Find the occupant/listing column and standardize it to LISTING.
    Handles:
    - LISTING column (ERIS format)
    - COMPANY_NAME column (other formats)
    - FACILITY_ID, OCCUPANT, TENANT, BUSINESS_NAME, etc.
    """
    cols_upper = {col.upper(): col for col in df.columns}
    
    # Pattern 1: LISTING already exists
    if "LISTING" in cols_upper:
        return df
    
    # Pattern 2: COMPANY_NAME (most common alternative)
    if "COMPANY_NAME" in cols_upper:
        df["LISTING"] = df[cols_upper["COMPANY_NAME"]]
        st.info(f"✓ Using {cols_upper['COMPANY_NAME']} as LISTING column")
        return df
    
    # Pattern 3: Other common patterns
    listing_patterns = [
        "FACILITY_ID",
        "OCCUPANT",
        "TENANT",
        "BUSINESS_NAME",
        "BUSINESS",
        "COMPANY",
        "NAME",
        "OCCUPANT_NAME"
    ]
    
    for pattern in listing_patterns:
        if pattern in cols_upper:
            df["LISTING"] = df[cols_upper[pattern]]
            st.info(f"✓ Using {cols_upper[pattern]} as LISTING column")
            return df
    
    return df

def read_input(file) -> pd.DataFrame:
    name = file.name.lower()

    if name.endswith(".csv"):
        df = pd.read_csv(file)
        df.columns = [str(c).strip().upper() for c in df.columns]
        if "ADDRESS" in df.columns:
            df["ADDRESS"] = df["ADDRESS"].ffill().apply(normalize_addr)
        return df

    # XLSX/XLS (requires openpyxl for .xlsx and xlrd for .xls)
    xls = pd.ExcelFile(file)
    raw = pd.read_excel(xls, sheet_name=0, header=None)

    header_row = None
    # First try to find ADDRESS + YEAR (ERIS format)
    for i in range(min(50, len(raw))):
        row_vals = raw.iloc[i].astype(str).str.upper().tolist()
        if "ADDRESS" in row_vals and "YEAR" in row_vals:
            header_row = i
            break
    
    # If not found, look for ADDRESS1 or COMPANY_NAME (other formats)
    if header_row is None:
        for i in range(min(50, len(raw))):
            row_vals = raw.iloc[i].astype(str).str.upper().tolist()
            if ("ADDRESS1" in row_vals or "COMPANY_NAME" in row_vals):
                header_row = i
                break

    if header_row is None:
        df = pd.read_excel(xls, sheet_name=0)
        df.columns = [str(c).strip().upper() for c in df.columns]
        if "ADDRESS" in df.columns:
            df["ADDRESS"] = df["ADDRESS"].ffill().apply(normalize_addr)
        return df

    df = pd.read_excel(xls, sheet_name=0, header=header_row)
    df.columns = [str(c).strip().upper() for c in df.columns]

    # Only filter by YEAR if the column exists
    if "YEAR" in df.columns:
        df = df[df["YEAR"].notna()]

    if "ADDRESS" in df.columns:
        df["ADDRESS"] = df["ADDRESS"].ffill().apply(normalize_addr)

    return df

def format_year_listing(df_addr: pd.DataFrame) -> pd.DataFrame:
    """Group by YEAR and combine listings into comma-separated unique string."""
    
    # Check if YEAR column exists
    has_year = "YEAR" in df_addr.columns
    
    if not has_year and "LISTING" not in df_addr.columns:
        return pd.DataFrame(columns=["Year(s)", "Occupant Listed"])
    
    # Format WITHOUT year data
    if not has_year:
        t = df_addr[["LISTING"]].copy()
        t["LISTING"] = t["LISTING"].astype(str).str.strip()
        t = t[t["LISTING"].str.len() > 0]
        
        # Remove duplicates
        unique_listings = t["LISTING"].drop_duplicates().tolist()
        
        result = pd.DataFrame({
            "Year(s)": ["N/A"] * len(unique_listings),
            "Occupant Listed": unique_listings
        })
        return result
    
    # Format WITH year data (original logic)
    if "LISTING" not in df_addr.columns:
        return pd.DataFrame(columns=["Year(s)", "Occupant Listed"])

    t = df_addr[["YEAR", "LISTING"]].copy()
    t["YEAR"] = pd.to_numeric(t["YEAR"], errors="coerce")
    t = t.dropna(subset=["YEAR"])
    t["YEAR"] = t["YEAR"].astype(int)

    t["LISTING"] = t["LISTING"].astype(str).str.strip()
    t = t[t["LISTING"].str.len() > 0]

    def combine_listings(series: pd.Series) -> str:
        seen = set()
        out = []
        for item in series.tolist():
            if item not in seen:
                seen.add(item)
                out.append(item)
        return ", ".join(out)

    grouped = (
        t.sort_values(["YEAR", "LISTING"], ascending=[True, True])
         .groupby("YEAR", as_index=False)["LISTING"]
         .apply(combine_listings)
         .rename(columns={"YEAR": "Year(s)", "LISTING": "Occupant Listed"})
         .reset_index(drop=True)
    )
    return grouped

def compress_year_runs(out_df: pd.DataFrame) -> list[tuple[str, str]]:
    """
    Turn:
      1970 A
      1971 A
      1972 B
    Into:
      1970-1971 A
      1972 B
    """
    if out_df.empty:
        return []

    years = out_df["Year(s)"].tolist()
    occs = out_df["Occupant Listed"].tolist()
    
    # If no year data (all "N/A"), just return the listings
    if all(str(y) == "N/A" for y in years):
        return [(y, occ) for y, occ in zip(years, occs)]

    rows: list[tuple[str, str]] = []
    start_y = years[0]
    prev_y = years[0]
    prev_occ = occs[0]

    for y, occ in zip(years[1:], occs[1:]):
        contiguous = (y == prev_y + 1)
        same_occ = (occ == prev_occ)
        if contiguous and same_occ:
            prev_y = y
            continue

        label = f"{start_y}-{prev_y}" if start_y != prev_y else f"{start_y}"
        rows.append((label, prev_occ))
        start_y = y
        prev_y = y
        prev_occ = occ

    label = f"{start_y}-{prev_y}" if start_y != prev_y else f"{start_y}"
    rows.append((label, prev_occ))
    return rows

def render_block(addr: str, kind: str, out_df: pd.DataFrame):
    st.markdown('<div class="addr-card">', unsafe_allow_html=True)
    st.markdown(
        f"""
        <div class="addr-header">
          City Directory Search for <span class="addr-pill">{addr}</span> ({kind})
        </div>
        """,
        unsafe_allow_html=True
    )

    rows_html = ""
    for _, r in out_df.iterrows():
        year = str(r.get("Year(s)", "")).strip()
        occ = str(r.get("Occupant Listed", "")).strip()
        rows_html += f"<tr><td>{year}</td><td>{occ}</td></tr>"

    table_html = f"""
      <table class="neat-table">
        <thead>
          <tr>
            <th>Year(s)</th>
            <th>Occupant Listed</th>
          </tr>
        </thead>
        <tbody>
          {rows_html if rows_html else "<tr><td colspan='2' style='opacity:.7;'>No results</td></tr>"}
        </tbody>
      </table>
    """
    st.markdown(table_html, unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

# ---------- DOCX helpers ----------
def set_cell_shading(cell, fill_hex: str):
    """fill_hex like 'D9EAD3' (light green)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_bold(cell, bold=True):
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = bold

def set_table_header_style(table, fill_hex="D9EAD3"):
    hdr = table.rows[0].cells
    for c in hdr:
        set_cell_shading(c, fill_hex)
        set_cell_bold(c, True)

def docx_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def build_subject_report_docx(subject_selected: list[str], df: pd.DataFrame) -> bytes:
    doc = Document()

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = True

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Year(s)"
    hdr_cells[1].text = "Subject Property Address(es) — Occupant Listed"
    set_table_header_style(table)

    for addr in subject_selected:
        block = df[df["ADDRESS"] == addr].copy()
        out = format_year_listing(block)
        runs = compress_year_runs(out)

        if not runs:
            row = table.add_row().cells
            row[0].text = ""
            row[1].text = f"{addr} — No results"
            continue

        for year_label, occ in runs:
            row = table.add_row().cells
            row[0].text = str(year_label)
            row[1].text = f"{addr} — {occ}"

    return docx_bytes(doc)

def build_adjoining_report_docx(adjoining_selected: list[str], df: pd.DataFrame, direction_map: dict) -> bytes:
    doc = Document()
    doc.add_paragraph("Addresses of adjoining properties were also reviewed. Historical tenants included:")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = True

    hdr = table.rows[0].cells
    hdr[0].text = "Direction"
    hdr[1].text = "Adjoining Property Addresses"
    hdr[2].text = "Occupant Listed (Year)"
    set_table_header_style(table)

    for addr in adjoining_selected:
        block = df[df["ADDRESS"] == addr].copy()
        out = format_year_listing(block)
        runs = compress_year_runs(out)

        direction = direction_map.get(addr, "")

        lines = []
        for year_label, occ in runs:
            if occ:
                lines.append(f"{occ} ({year_label})")
        occ_text = "\n".join(lines) if lines else "No results"

        row = table.add_row().cells
        row[0].text = direction
        row[1].text = addr
        row[2].text = occ_text

    return docx_bytes(doc)

# ---------- Upload ----------
uploaded = st.file_uploader("Upload City Directory export (CSV, XLSX, or XLS)", type=["csv", "xlsx", "xls"])
if not uploaded:
    st.stop()

df = read_input(uploaded)
df.columns = [c.upper() for c in df.columns]

# Find and combine address columns
df = find_and_combine_address_columns(df)

if "ADDRESS" not in df.columns:
    st.error("❌ Could not find an ADDRESS column. Available columns: " + ", ".join(df.columns))
    st.info("Looking for: ADDRESS, ADDRESS1/ADDRESS2, STREET ADDRESS, PROPERTY ADDRESS, etc.")
    st.stop()

# Find and standardize listing/occupant column
df = find_listing_column(df)

if "LISTING" not in df.columns:
    st.error("❌ Could not find a LISTING/COMPANY_NAME column. Available columns: " + ", ".join(df.columns))
    st.info("Looking for: LISTING, COMPANY_NAME, FACILITY_ID, OCCUPANT, TENANT, etc.")
    st.stop()

# Check if YEAR column exists and warn user
if "YEAR" not in df.columns:
    st.warning("⚠️ No YEAR column found in this file. Results will show occupants without year information.")

# ✅ UPDATED SORTING: street name alpha, then house number numeric, then unit numeric
all_addresses = [a for a in df["ADDRESS"].dropna().unique() if str(a).strip()]
all_addresses = sorted(all_addresses, key=parse_address_for_sort)

st.success(f"Loaded {len(df):,} rows • Found {len(all_addresses):,} unique addresses")

# ---------- Session state + callbacks ----------
if "subject_sel" not in st.session_state:
    st.session_state["subject_sel"] = []
if "adjoining_sel" not in st.session_state:
    st.session_state["adjoining_sel"] = []
if "run_subject" not in st.session_state:
    st.session_state["run_subject"] = False
if "run_adjoining" not in st.session_state:
    st.session_state["run_adjoining"] = False
if "dir_map" not in st.session_state:
    st.session_state["dir_map"] = {}

def clear_all():
    st.session_state["subject_sel"] = []
    st.session_state["adjoining_sel"] = []
    st.session_state["run_subject"] = False
    st.session_state["run_adjoining"] = False
    st.session_state["dir_map"] = {}

def set_run_subject():
    st.session_state["run_subject"] = True

def set_run_adjoining():
    st.session_state["run_adjoining"] = True

# ---------- TOP UI ----------
ui_left, ui_right = st.columns(2)

with ui_left:
    st.subheader("Pick Subject Property Addresses")
    subject_selected = st.multiselect("Subject addresses", all_addresses, key="subject_sel", placeholder="Choose address")
    st.button("CREATE SUBJECT PROPERTY TABLES", use_container_width=True, on_click=set_run_subject)

with ui_right:
    st.subheader("Pick Adjoining Property Addresses")
    adjoining_selected = st.multiselect("Adjoining addresses", all_addresses, key="adjoining_sel", placeholder="Choose address")
    st.button("CREATE ADJOINING PROPERTY TABLES", use_container_width=True, on_click=set_run_adjoining)

st.button("CLEAR ALL", use_container_width=True, on_click=clear_all)
st.divider()

# ---------- OUTPUT IN TWO COLUMNS ----------
out_left, out_right = st.columns(2)

with out_left:
    st.markdown('<h2 class="section-title">Subject Property Tables</h2>', unsafe_allow_html=True)

    if st.session_state["run_subject"] and subject_selected:
        for addr in subject_selected:
            block = df[df["ADDRESS"] == addr].copy()
            out = format_year_listing(block)
            render_block(addr, "Subject Property", out)

        subj_docx = build_subject_report_docx(subject_selected, df)
        st.download_button(
            "Download Subject Report Table (.docx)",
            data=subj_docx,
            file_name="ELC_Subject_Report_Table.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    else:
        st.caption("Select a subject address and click CREATE SUBJECT PROPERTY TABLES.")

with out_right:
    st.markdown('<h2 class="section-title">Adjoining Property Tables</h2>', unsafe_allow_html=True)

    if st.session_state["run_adjoining"] and adjoining_selected:

        with st.expander("Optional: Set directions for adjoining addresses (North/East/South/West)", expanded=False):
            dir_opts = ["", "North", "East", "South", "West"]
            for a in adjoining_selected:
                key = f"dir_{a}"
                if a not in st.session_state["dir_map"]:
                    st.session_state["dir_map"][a] = ""
                picked = st.selectbox(a, dir_opts, index=dir_opts.index(st.session_state["dir_map"][a]), key=key)
                st.session_state["dir_map"][a] = picked

        scroll_box = st.container(height=720)
        with scroll_box:
            for addr in adjoining_selected:
                block = df[df["ADDRESS"] == addr].copy()
                out = format_year_listing(block)
                render_block(addr, "Adjoining Property", out)

        adj_docx = build_adjoining_report_docx(adjoining_selected, df, st.session_state["dir_map"])
        st.download_button(
            "Download Adjoining Report Table (.docx)",
            data=adj_docx,
            file_name="ELC_Adjoining_Report_Table.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    else:
        st.caption("Select adjoining addresses and click CREATE ADJOINING PROPERTY TABLES.")